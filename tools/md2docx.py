# -*- coding: utf-8 -*-
"""将Markdown转换为单栏DOCX（通用），含LaTeX公式渲染为Word原生公式"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import subprocess
import tempfile
import copy
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# 字号常量（全部小一号）
SIZE_XIAO_ER = Pt(16)   # 三号（原小二）
SIZE_SAN_HAO = Pt(14)   # 四号（原三号）
SIZE_SI_HAO = Pt(12)    # 小四（原四号）
SIZE_XIAO_SI = Pt(10.5) # 五号（原小四）

# 版面参数。可在调用 process(..., layout={...}) 时覆盖，用于精调页数。
PAGE_MARGIN_CM = 2
IMAGE_WIDTH_CM = 8
BODY_SPACE_AFTER_PT = 3
OPTION_SPACE_AFTER_PT = 2
IMAGE_SPACE_AFTER_PT = 4
H1_SPACE_BEFORE_PT = 12
H1_SPACE_AFTER_PT = 8
H2_SPACE_BEFORE_PT = 4
H2_SPACE_AFTER_PT = 8

# 字体名常量
FONT_HEI = '黑体'
FONT_FANGSONG = '仿宋'
FONT_FZXBS = '方正小标宋简体'
FONT_TIMES = 'Times New Roman'

# LaTeX 符号映射到 Unicode
LATEX_TO_UNICODE = {
    r'\prime': '\u2032',  # ′
    r'\circ': '\u00B0',   # °
    r'\triangle': '\u25B3',  # △
    r'\angle': '\u2220',     # ∠
    r'\parallel': '\u2225',  # ∥
    r'\perp': '\u27C2',      # ⟂
    r'\cdot': '\u00B7',     # ·
    r'\times': '\u00D7',    # ×
    r'\div': '\u00F7',      # ÷
    r'\pm': '\u00B1',       # ±
    r'\le': '\u2264',       # ≤
    r'\ge': '\u2265',       # ≥
    r'\ne': '\u2260',       # ≠
    r'\to': '\u2192',       # →
    r'\Rightarrow': '\u21D2',  # ⇒
    r'\pi': '\u03C0',       # π
    r'\alpha': '\u03B1',
    r'\beta': '\u03B2',
    r'\gamma': '\u03B3',
    r'\delta': '\u03B4',
    r'\theta': '\u03B8',
    r'\vert': '|',          # |
    r'\scriptstyle': '',    # 删除（仅影响字号）
    r'\frac': None,
    r'\sqrt': None,
}

def set_columns(section, num=1):
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:cols')):
        sectPr.remove(old)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num))
    if num == 2:
        cols.set(qn('w:space'), '425')
        cols.set(qn('w:equalWidth'), '1')
        cols.set(qn('w:sep'), '1')
    sectPr.append(cols)

def add_page_number(section, start=None):
    """为分节添加底部居中页码
    start: 起始页码，None 表示延续上一节的编号"""
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    sectPr.insert(0, pgNumType)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run1 = p.add_run()
    run1._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2 = p.add_run()
    run2._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3 = p.add_run()
    run3._r.append(fldChar_end)

    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = FONT_TIMES

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)

def set_vert_align(run, val):
    """设置上标/下标"""
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), val)
    rPr.append(vertAlign)

def set_run_font(run, cn_font, en_font, size, italic_en=True, color=RGBColor(0, 0, 0)):
    """设置中英文字体、字号、颜色"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), en_font)
    run.font.size = size
    run.font.color.rgb = color
    # 设置italic - 通过 XML
    if italic_en:
        i = OxmlElement('w:i')
        i.set(qn('w:val'), '1')
        # 只对西文生效需要复杂控制，这里简化：英文倾斜，中文也倾斜
        # 如需区分，需要更精细的字符级处理
        # 此处保持整体 italic 即可（Word中会同时影响中英文）

def split_cn_en(text):
    """分离中文与英文/数字（用于不同字体设置）
    返回: [(text, is_latin)]"""
    if not text:
        return []
    result = []
    current = ''
    current_is_latin = None
    for ch in text:
        is_latin = (ord(ch) < 0x2E80)  # 简化判断：基本拉丁及扩展
        if current_is_latin is None:
            current = ch
            current_is_latin = is_latin
        elif is_latin == current_is_latin:
            current += ch
        else:
            result.append((current, current_is_latin))
            current = ch
            current_is_latin = is_latin
    if current:
        result.append((current, current_is_latin))
    return result

def add_styled_runs(p, text, cn_font, en_font, size, italic_en=True):
    """添加带中英字体分离的runs"""
    if not text:
        return
    segments = split_cn_en(text)
    for seg_text, is_latin in segments:
        run = p.add_run(seg_text)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        # 英文段：英文字体；中文段：中文字体
        font = en_font if is_latin else cn_font
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rFonts.set(qn('w:eastAsia'), font)
        rFonts.set(qn('w:cs'), font)
        run.font.size = size
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 字母倾斜（仅英文段）
        if italic_en and is_latin:
            i = OxmlElement('w:i')
            i.set(qn('w:val'), '1')
            rPr.append(i)

def apply_layout(layout):
    """Override selected layout globals for one export run."""
    if not layout:
        return
    global SIZE_XIAO_ER, SIZE_SAN_HAO, SIZE_SI_HAO, SIZE_XIAO_SI
    global PAGE_MARGIN_CM, IMAGE_WIDTH_CM
    global BODY_SPACE_AFTER_PT, OPTION_SPACE_AFTER_PT, IMAGE_SPACE_AFTER_PT
    global H1_SPACE_BEFORE_PT, H1_SPACE_AFTER_PT, H2_SPACE_BEFORE_PT, H2_SPACE_AFTER_PT

    if 'h1_pt' in layout:
        SIZE_XIAO_ER = Pt(layout['h1_pt'])
    if 'h2_pt' in layout:
        SIZE_SAN_HAO = Pt(layout['h2_pt'])
    if 'body_pt' in layout:
        SIZE_SI_HAO = Pt(layout['body_pt'])
    if 'table_pt' in layout:
        SIZE_XIAO_SI = Pt(layout['table_pt'])
    if 'margin_cm' in layout:
        PAGE_MARGIN_CM = layout['margin_cm']
    if 'image_width_cm' in layout:
        IMAGE_WIDTH_CM = layout['image_width_cm']
    if 'body_space_after_pt' in layout:
        BODY_SPACE_AFTER_PT = layout['body_space_after_pt']
    if 'option_space_after_pt' in layout:
        OPTION_SPACE_AFTER_PT = layout['option_space_after_pt']
    if 'image_space_after_pt' in layout:
        IMAGE_SPACE_AFTER_PT = layout['image_space_after_pt']
    if 'h1_space_before_pt' in layout:
        H1_SPACE_BEFORE_PT = layout['h1_space_before_pt']
    if 'h1_space_after_pt' in layout:
        H1_SPACE_AFTER_PT = layout['h1_space_after_pt']
    if 'h2_space_before_pt' in layout:
        H2_SPACE_BEFORE_PT = layout['h2_space_before_pt']
    if 'h2_space_after_pt' in layout:
        H2_SPACE_AFTER_PT = layout['h2_space_after_pt']

def add_paragraph_with_math(doc, text, cn_font=FONT_FANGSONG, en_font=FONT_TIMES,
                              size=SIZE_SI_HAO, left_indent=None, space_after=3,
                              italic_en=True):
    """添加包含LaTeX公式的段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    p.paragraph_format.space_after = Pt(space_after)

    pattern = re.compile(r'\$([^$]+)\$')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            plain = text[pos:m.start()]
            add_styled_runs(p, plain, cn_font, en_font, size, italic_en)
        latex = m.group(1).strip()
        process_latex_with_supsub(p, latex, cn_font, en_font, size, italic_en)
        pos = m.end()
    if pos < len(text):
        plain = text[pos:]
        add_styled_runs(p, plain, cn_font, en_font, size, italic_en)
    return p

def process_latex_with_supsub(p, latex, cn_font, en_font, base_size, italic_en=True):
    """处理含上下标的LaTeX"""
    latex = re.sub(r'\\frac\s*\{([^}]*)\}\s*\{([^}]*)\}',
                   lambda m: f'({m.group(1).strip()})/({m.group(2).strip()})', latex)
    # 先去掉所有花括号内的多余空格，避免干扰后续处理
    latex = re.sub(r'\{\s*([^{}]*?)\s*\}', r'{\1}', latex)
    for sym, uni in LATEX_TO_UNICODE.items():
        if uni is None:
            continue
        latex = re.sub(re.escape(sym), uni, latex)
    latex = re.sub(r'\\sqrt\s*\{([^}]*)\}', lambda m: f'√({m.group(1).strip()})', latex)
    latex = re.sub(r'\\left\s*([\(\[\|])', r'\1', latex)
    latex = re.sub(r'\\right\s*([\)\]\|])', r'\1', latex)
    latex = re.sub(r'\\displaystyle', '', latex)
    latex = re.sub(r'\\text\s*\{([^}]*)\}', r'\1', latex)
    latex = re.sub(r'\\mathrm\s*\{([^}]*)\}', r'\1', latex)
    latex = re.sub(r'\\,|\\;|\\:|\\!', '', latex)

    tokens = re.split(r'(\^[^{]?[^{]*?\{[^}]*\}|_[^{]?[^{]*?\{[^}]*\}|\^[a-zA-Z0-9]|_[a-zA-Z0-9])', latex)

    sup_size = Pt(int(base_size.pt * 0.7))
    for token in tokens:
        if not token:
            continue
        if token.startswith('^'):
            sup = token[1:].strip('{}').strip()
            sup = re.sub(r'\s+', '', sup)  # 去除内部多余空格
            sup = re.sub(r'\\prime', '\u2032', sup)
            sup = re.sub(r'\\circ', '\u00B0', sup)
            sup = re.sub(r'[{}]', '', sup)  # 去除上标内的花括号
            sup = sup.strip()  # 去除首尾空格
            add_styled_runs(p, sup, cn_font, en_font, sup_size, italic_en)
            # 设置最后一个run为上标
            for r in p.runs[::-1]:
                if r.text.strip() == sup or r.text.strip().endswith(sup):
                    set_vert_align(r, 'superscript')
                    break
        elif token.startswith('_'):
            sub = token[1:].strip('{}').strip()
            sub = re.sub(r'\s+', '', sub)  # 去除内部多余空格
            sub = re.sub(r'[{}]', '', sub)  # 去除下标内的花括号
            sub = sub.strip()  # 去除首尾空格
            add_styled_runs(p, sub, cn_font, en_font, sup_size, italic_en)
            for r in p.runs[::-1]:
                if r.text.strip() == sub or r.text.strip().endswith(sub):
                    set_vert_align(r, 'subscript')
                    break
        else:
            # 清理普通文本中的花括号
            token = re.sub(r'[{}]', '', token)
            token = re.sub(r'\s+', ' ', token)  # 合并多余空格
            add_styled_runs(p, token, cn_font, en_font, base_size, italic_en)

def _find_pandoc():
    pandoc_exe = 'pandoc'
    try:
        subprocess.run([pandoc_exe, '--version'], capture_output=True, check=True)
        return pandoc_exe
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass
    # 尝试常见安装路径
    for path in [
        os.path.expandvars(r'%LOCALAPPDATA%\Pandoc\pandoc.exe'),
        os.path.expandvars(r'%PROGRAMFILES%\Pandoc\pandoc.exe'),
        r'C:\Program Files\Pandoc\pandoc.exe',
    ]:
        if os.path.exists(path):
            return path
    # 搜索 winget 安装路径
    winget_base = os.path.expandvars(r'%LOCALAPPDATA%\Microsoft\WinGet\Packages')
    if os.path.isdir(winget_base):
        for root, dirs, files in os.walk(winget_base):
            for f in files:
                if f == 'pandoc.exe':
                    return os.path.join(root, f)
            if root.count(os.sep) > winget_base.count(os.sep) + 3:
                dirs.clear()
    raise RuntimeError('未找到 Pandoc，请安装后重试')

def _strip_yaml(content):
    return re.sub(r'^---\s*\n.*?\n---\s*\n', '', content, count=1, flags=re.DOTALL)

def _is_variable_run(text):
    """判断 run 是否为数学变量（应倾斜）：
    纯拉丁字母+数字+标点，至少一个字母，不含中文文字"""
    stripped = text.strip()
    if not stripped:
        return False
    # 不能包含中文字（排除文字 run），但允许中文标点
    if re.search(r'[\u4e00-\u9fff]', stripped):
        return False
    # 必须至少有一个拉丁字母
    if not re.search(r'[A-Za-z]', stripped):
        return False
    return True

def _modify_paragraph_style(para, cn_font, en_font, size, bold=False):
    """直接修改段落中所有 run 的字体，变量自动倾斜"""
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:cs'), en_font)
        run.font.size = size
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = bold
        # 数学变量（单独拉丁字母）倾斜
        if _is_variable_run(run.text):
            i = OxmlElement('w:i')
            i.set(qn('w:val'), '1')
            rPr.append(i)

def _insert_separator_before(para):
    """在段落前插入一个带底部边框的空段落（分隔线）"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(6)

def _is_option_line(text):
    return bool(re.match(r'^[A-D][．.]', text))

def _is_inline_table_row(text):
    return text.startswith('|')

def _adjust_images(doc):
    """调整非表格段落中图片宽度为统一值，保持横纵比；表格内图片不处理"""
    for para in doc.paragraphs:
        # 跳过表格内段落
        parent = para._p.getparent()
        parent_tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
        if parent_tag == 'tc':
            continue
        for run in para.runs:
            if run._r.findall(qn('w:drawing')) or run._r.findall(qn('w:pict')):
                for inline in para._p.findall('.//' + qn('wp:inline')):
                    extent = inline.find(qn('wp:extent'))
                    if extent is not None:
                        old_cx = extent.get('cx')
                        old_cy = extent.get('cy')
                        if old_cx and old_cy:
                            old_w = int(old_cx)
                            old_h = int(old_cy)
                            new_w = int(IMAGE_WIDTH_CM * 360000)
                            new_h = int(new_w * old_h / old_w) if old_w > 0 else new_w
                            extent.set('cx', str(new_w))
                            extent.set('cy', str(new_h))
                for anchor in para._p.findall('.//' + qn('wp:anchor')):
                    extent = anchor.find(qn('wp:extent'))
                    if extent is not None:
                        old_cx = extent.get('cx')
                        old_cy = extent.get('cy')
                        if old_cx and old_cy:
                            old_w = int(old_cx)
                            old_h = int(old_cy)
                            new_w = int(IMAGE_WIDTH_CM * 360000)
                            new_h = int(new_w * old_h / old_w) if old_w > 0 else new_w
                            extent.set('cx', str(new_w))
                            extent.set('cy', str(new_h))

OPTION_IMG_TABLE_CM = 5.5  # 选项图片表格总宽度

def _rearrange_option_images(doc):
    """2张以上连续图片自动入表（有标签优先，无标签也处理）"""
    body = doc.element.body
    paras = doc.paragraphs
    i = 0
    while i < len(paras):
        group, consumed = _find_image_group(paras, i)
        if group and len(group) >= 2:
            labels = [g['label'] for g in group]
            images = [g['drawing'] for g in group]
            first_idx = _body_index(body, group[0]['para']._p)
            old_count = consumed
            tbl = _build_image_option_table(labels, images, OPTION_IMG_TABLE_CM)
            body.insert(first_idx, tbl)
            for offset in range(old_count - 1, -1, -1):
                body.remove(paras[i + offset]._p)
            paras = doc.paragraphs
            continue
        i += consumed if consumed else 1

def _find_image_group(paras, i):
    """从 i 开始找连续图片组，返回 ([{label,drawing,para}], consumed_count) 或 (None, 0)"""
    group = []
    j = i
    while j < len(paras):
        has_img = _para_has_image(paras[j])
        if has_img:
            drawing = _extract_drawing(paras[j])
            if drawing is None:
                break
            # 检查下一段是否是单字母标签
            label = None
            if j + 1 < len(paras):
                next_text = paras[j + 1].text.strip()
                if _is_single_letter_pos(next_text) is not None:
                    label = next_text
            group.append({'label': label, 'drawing': drawing, 'para': paras[j]})
            j += 1
            if label is not None:
                j += 1  # 跳过标签段
            continue
        # 纯图（无标签间隔）：图+图+图
        if group and not has_img:
            break
        break
    if len(group) < 2:
        return None, 0
    consumed = 0
    for g in group:
        consumed += 1  # 图片段
        if g['label'] is not None:
            consumed += 1  # 标签段
    return group, consumed

def _is_single_letter_pos(text):
    for ch in 'ABCD':
        if re.match(r'^' + ch + r'\s*$', text):
            return ch
    return None

def _build_image_option_table(labels, images, total_width_cm):
    """构建图片选项表格：1行N列，含标签（若有）"""
    n = len(images)
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblWidth')
    tblW.set(qn('w:w'), str(int(total_width_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{bn}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tblGrid = OxmlElement('w:tblGrid')
    cell_w = int(total_width_cm / n * 567)
    for _ in range(n):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(cell_w))
        tblGrid.append(gridCol)
    tblPr.append(tblGrid)
    tbl.append(tblPr)

    cell_w_cm = total_width_cm / n
    tr = OxmlElement('w:tr')
    for idx in range(n):
        tc = _build_option_cell(labels[idx] if labels[idx] else '', images[idx], cell_w_cm)
        tr.append(tc)
    tbl.append(tr)
    return tbl

def _build_option_cell(label, drawing, cell_width_cm):
    tc = OxmlElement('w:tc')
    if label:
        # 字母
        p_lbl = OxmlElement('w:p')
        pPr_lbl = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr_lbl.append(jc)
        sp_after = OxmlElement('w:spacing')
        sp_after.set(qn('w:after'), '40')
        pPr_lbl.append(sp_after)
        p_lbl.append(pPr_lbl)
        r_lbl = OxmlElement('w:r')
        rPr_lbl = OxmlElement('w:rPr')
        rFonts_lbl = OxmlElement('w:rFonts')
        rFonts_lbl.set(qn('w:ascii'), FONT_TIMES)
        rFonts_lbl.set(qn('w:hAnsi'), FONT_TIMES)
        rFonts_lbl.set(qn('w:eastAsia'), FONT_TIMES)
        rPr_lbl.append(rFonts_lbl)
        sz_lbl = OxmlElement('w:sz')
        sz_lbl.set(qn('w:val'), str(int(SIZE_XIAO_SI.pt * 2)))
        rPr_lbl.append(sz_lbl)
        b_lbl = OxmlElement('w:b')
        rPr_lbl.append(b_lbl)
        r_lbl.append(rPr_lbl)
        t_lbl = OxmlElement('w:t')
        t_lbl.set(qn('xml:space'), 'preserve')
        t_lbl.text = label
        r_lbl.append(t_lbl)
        p_lbl.append(r_lbl)
        tc.append(p_lbl)
    # 图片，缩放到单元格宽度
    drawing_copy = copy.deepcopy(drawing)
    _scale_drawing_to_width(drawing_copy, cell_width_cm)
    p_img = OxmlElement('w:p')
    pPr_img = OxmlElement('w:pPr')
    jc_img = OxmlElement('w:jc')
    jc_img.set(qn('w:val'), 'center')
    pPr_img.append(jc_img)
    p_img.append(pPr_img)
    r_img = OxmlElement('w:r')
    r_img.append(drawing_copy)
    p_img.append(r_img)
    tc.append(p_img)
    return tc

def _scale_drawing_to_width(drawing, target_width_cm):
    """缩放 drawing 元素的宽高到目标宽度，保持横纵比"""
    for extent in drawing.findall('.//' + qn('wp:extent')):
        cx = extent.get('cx')
        if cx:
            old_w = int(cx)
            old_cy = extent.get('cy')
            old_h = int(old_cy) if old_cy else old_w
            new_w = int(target_width_cm * 360000)
            new_h = int(new_w * old_h / old_w) if old_w > 0 else new_w
            extent.set('cx', str(new_w))
            extent.set('cy', str(new_h))

def _para_has_image(para):
    return bool(para._p.findall('.//' + qn('wp:inline'))) or bool(para._p.findall('.//' + qn('wp:anchor')))

def _para_has_omml(para):
    return bool(para._p.findall('.//' + qn('m:oMath')))

def _is_single_letter(text, expected):
    return bool(re.match(r'^' + expected + r'\s*$', text.strip()))

def _body_index(body, element):
    for idx, child in enumerate(body):
        if child is element:
            return idx
    raise ValueError('element not found in body')

def _extract_drawing(para):
    for elem in para._p.findall('.//' + qn('w:drawing')):
        return elem
    return None

def _parse_md_table(lines):
    """解析MD表格行，返回 (headers, data_rows)"""
    headers = [c.strip() for c in lines[0].split('|')[1:-1]]
    data = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            data.append(cells)
    return headers, data

def _find_md_tables(content):
    """查找MD中的表格块，返回 [(start, end, lines)] 位置"""
    lines = content.split('\n')
    tables = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith('|') and '|' in stripped[1:]:
            j = i
            table_lines = []
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1
            if len(table_lines) >= 3:
                sep = table_lines[1]
                if re.match(r'^\|[\s\-:|]+\|$', sep):
                    tables.append((i, j, table_lines))
            i = j
        else:
            i += 1
    return tables

def _render_table_to_image(headers, data, output_path, cn_font_size=9):
    """用matplotlib渲染表格为PNG图片，宽度限制8cm"""
    n_rows = len(data) + 1
    n_cols = len(headers)
    cell_text = [headers] + data

    max_width_inch = 8 / 2.54

    # 根据列数动态调整字体和行高
    if n_cols <= 4:
        font_size = cn_font_size
        row_height_factor = 1.4
        cell_width_inch = max_width_inch / n_cols
    elif n_cols <= 7:
        font_size = cn_font_size - 1
        row_height_factor = 1.3
        cell_width_inch = max_width_inch / n_cols
    else:
        font_size = cn_font_size - 2
        row_height_factor = 1.2
        cell_width_inch = max_width_inch / n_cols

    fig_w = max_width_inch
    fig_h = max(n_rows * 0.4, 1.2)

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis('off')
    ax.set_xlim(0, max_width_inch)
    ax.set_ylim(0, n_rows * 0.4)

    tbl = ax.table(
        cellText=cell_text,
        cellLoc='center',
        loc='center',
        cellColours=None,
    )

    tbl.auto_set_font_size(False)
    tbl.scale(1, row_height_factor)

    for (row, col), cell in tbl.get_celld().items():
        cell.set_edgecolor('#333333')
        cell.set_linewidth(0.5)
        cell.set_fontsize(font_size)
        if row == 0:
            cell.set_facecolor('#e8e8e8')
            cell.set_text_props(weight='bold', fontsize=font_size)
        else:
            cell.set_facecolor('#ffffff')
            cell.set_text_props(fontsize=font_size)

    plt.tight_layout(pad=0.3)
    fig.savefig(output_path, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return output_path

def _setup_mpl_chinese():
    """设置 matplotlib 中文字体"""
    for font_name in ['Microsoft YaHei', 'SimHei', 'STSong']:
        try:
            from matplotlib import font_manager
            if font_manager.findfont(font_name):
                plt.rcParams['font.sans-serif'] = [font_name, 'DejaVu Sans']
                plt.rcParams['axes.unicode_minus'] = False
                return font_name
        except Exception:
            continue
    return None

def _clean_cell_text(text):
    """清理单元格文本：去 $、处理上下标花括号、去多余空格"""
    text = re.sub(r'\$', '', text)
    text = re.sub(r'\^\{([^}]*)\}', r'^\1', text)
    text = re.sub(r'\_\{([^}]*)\}', r'_\1', text)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def _replace_tables_with_images(content, md_dir):
    """将MD中的表格替换为图片引用，返回新content和图片路径列表"""
    _setup_mpl_chinese()

    tables = _find_md_tables(content)
    if not tables:
        return content, []

    lines = content.split('\n')
    image_paths = []

    for start, end, table_lines in reversed(tables):
        headers, data = _parse_md_table(table_lines)
        headers = [_clean_cell_text(h) for h in headers]
        data = [[_clean_cell_text(c) for c in row] for row in data]
        img_path = os.path.join(md_dir, f'_table_img_{start}.png')
        _render_table_to_image(headers, data, img_path)
        image_paths.append(img_path)
        img_rel_path = os.path.basename(img_path)
        replacement = f'![]({img_rel_path})'
        lines[start:end] = [replacement, '']

    return '\n'.join(lines), image_paths


_TEMP_IMG_DIR = None

def _cleanup_table_images():
    global _TEMP_IMG_DIR
    if _TEMP_IMG_DIR and os.path.isdir(_TEMP_IMG_DIR):
        import shutil
        shutil.rmtree(_TEMP_IMG_DIR, ignore_errors=True)
        _TEMP_IMG_DIR = None

_TEXT_OPTION_TABLE_CM = 7.5   # 文字选项总宽度

def _rearrange_text_options(doc):
    """A/B/C/D 文字选项重组段落：优先合并为单行，次之两行，不用表格"""
    body = doc.element.body
    paras = doc.paragraphs
    i = 0
    while i < len(paras):
        if i + 3 >= len(paras):
            break
        if not _is_consecutive_text_options(paras, i):
            i += 1
            continue
        items = []
        for j in range(4):
            p_elem = paras[i + j]._p
            text = _extract_para_text(p_elem)
            has_omml = _para_has_omml(paras[i + j])
            items.append({'text': text.strip(), 'p_elem': p_elem, 'has_omml': has_omml})
        texts = [it['text'] for it in items]
        layout = _decide_text_option_layout(texts)
        # 含公式的选项强制双行，避免宽度估算不准导致挤压
        if layout and layout[0] == 1 and any(it['has_omml'] for it in items):
            if max([_estimate_text_width_cm(t) for t in texts]) <= _TEXT_OPTION_TABLE_CM / 2 * 0.95:
                layout = (2, 2)
        if layout is None:
            i += 4
            continue
        rows, cols = layout
        # 在移除原段落前，先捕获所有子元素（避免detached元素问题）
        for item in items:
            if item['has_omml']:
                item['_children_xml'] = []
                from lxml import etree
                for child in item['p_elem']:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'pPr':
                        continue
                    item['_children_xml'].append(etree.tostring(child, encoding='unicode'))
            else:
                item['_children_xml'] = None
        first_idx = _body_index(body, items[0]['p_elem'])
        old_elems = [paras[i + j]._p for j in range(4)]
        for elem in reversed(old_elems):
            body.remove(elem)
        if rows == 1:
            new_p = _merge_options_to_paragraph(items, cols)
            body.insert(first_idx, new_p)
        else:
            p1 = _merge_options_to_paragraph(items[:2], 2)
            p2 = _merge_options_to_paragraph(items[2:], 2)
            body.insert(first_idx, p1)
            body.insert(first_idx + 1, p2)
        paras = doc.paragraphs

def _merge_options_to_paragraph(items, cols):
    """合并选项到一个段落，用tab分隔，设置tab stops控制对齐"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)
    # 设置tab stops
    tabs = OxmlElement('w:tabs')
    cell_w_dxa = int(_TEXT_OPTION_TABLE_CM / cols * 567)
    for k in range(1, cols):
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(k * cell_w_dxa))
        tabs.append(tab)
    pPr.append(tabs)
    sp_after = OxmlElement('w:spacing')
    sp_after.set(qn('w:after'), str(OPTION_SPACE_AFTER_PT * 20))
    pPr.append(sp_after)
    p.append(pPr)

    for idx, item in enumerate(items):
        if idx > 0:
            r_tab = OxmlElement('w:r')
            tab_elem = OxmlElement('w:tab')
            r_tab.append(tab_elem)
            p.append(r_tab)
        if item['has_omml']:
            from lxml import etree
            for xml_str in item['_children_xml']:
                child = etree.fromstring(xml_str)
                p.append(child)
        else:
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), FONT_TIMES)
            rFonts.set(qn('w:hAnsi'), FONT_TIMES)
            rFonts.set(qn('w:eastAsia'), FONT_FANGSONG)
            rFonts.set(qn('w:cs'), FONT_TIMES)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(SIZE_SI_HAO.pt * 2)))
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = item['text']
            r.append(t)
            p.append(r)
    return p

def _extract_para_text(p_elem):
    """提取段落纯文本（含 OMML 公式），OMML 部分用空白代替"""
    parts = []
    for r in p_elem.findall(qn('w:r')):
        is_omml = bool(r.findall(qn('m:oMath')) or r.findall('.//' + qn('m:oMath')))
        if is_omml:
            parts.append(' ')
            continue
        for t in r.findall(qn('w:t')):
            if t.text:
                parts.append(t.text)
    # 直接挂在p下的OMML也算
    if p_elem.findall(qn('m:oMath')):
        parts.append(' ')
    return ''.join(parts)

def _is_consecutive_text_options(paras, i):
    for j, letter in enumerate('ABCD'):
        if i + j >= len(paras):
            return False
        text = _extract_para_text(paras[i + j]._p)
        if not re.match(r'^' + letter + r'[．.]', text.strip()):
            return False
    return True

def _estimate_text_width_cm(text):
    """粗略估算文字在 小四(12pt) 下的宽度(cm)，偏保守（略高估）"""
    w = 0.0
    for ch in text:
        if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef':
            w += 0.47  # 中文/全角 ~12pt 仿宋
        elif ord(ch) > 0x7f:
            w += 0.47
        else:
            w += 0.24  # 英文/数字/符号
    w += 0.05  # 表格内边距补偿
    return w

def _decide_text_option_layout(texts):
    """决定1行或2x2或放弃
    返回 (rows, cols) 或 None"""
    widths = [_estimate_text_width_cm(t) for t in texts]
    total_w = sum(widths)
    cell_w_1row = _TEXT_OPTION_TABLE_CM / 4
    cell_w_2row = _TEXT_OPTION_TABLE_CM / 2
    # 优先单行：每个都能放进1/4格
    if max(widths) <= cell_w_1row * 0.98:
        return (1, 4)
    # 次之双行：每个都能放进半宽单元格
    if max(widths) <= cell_w_2row * 0.95:
        return (2, 2)
    # 不行就保留原样
    return None

def process(md_file, docx_file, layout=None):
    """使用 Pandoc 处理公式 + python-docx 控制排版"""
    apply_layout(layout)

    pandoc_exe = _find_pandoc()
    md_dir = os.path.dirname(os.path.abspath(md_file))

    # 读取 MD，去掉 YAML 元信息
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    content = _strip_yaml(content)

    # 将表格替换为图片
    main_content, table_images = _replace_tables_with_images(content, md_dir)

    # 写临时 MD 给 Pandoc
    tmp_md = tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8')
    tmp_md.write(main_content)
    tmp_md.close()

    tmp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name

    try:
        subprocess.run(
            [pandoc_exe, tmp_md.name, '-o', tmp_docx, '--from', 'markdown', '--to', 'docx'],
            cwd=md_dir, check=True, capture_output=True, text=True,
        )
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f'Pandoc 转换失败:\n{e.stderr}') from e
    finally:
        os.unlink(tmp_md.name)

    # 打开 Pandoc 生成的 DOCX，套用排版
    doc = Document(tmp_docx)

    # 选项图片重组为表格
    _rearrange_option_images(doc)
    # 文字选项重组段落
    _rearrange_text_options(doc)

    # 设置页面：A4 / 2cm 边距 / 单栏
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)
    set_columns(section, num=1)
    add_page_number(section, start=1)

    # 逐段落调整字体和排版
    for para in doc.paragraphs:
        style_name = (para.style.name or '').lower()

        if 'heading 2' in style_name:
            # H2：方正小标宋 四号 加粗 + 前插分隔线
            _insert_separator_before(para)
            _modify_paragraph_style(para, FONT_FZXBS, FONT_TIMES, SIZE_SAN_HAO, bold=True)
            para.paragraph_format.space_after = Pt(H2_SPACE_AFTER_PT)
            para.paragraph_format.space_before = Pt(H2_SPACE_BEFORE_PT)

        elif 'heading 1' in style_name:
            # H1：黑体 三号 加粗
            _modify_paragraph_style(para, FONT_HEI, FONT_TIMES, SIZE_XIAO_ER, bold=True)
            para.paragraph_format.space_after = Pt(H1_SPACE_AFTER_PT)
            para.paragraph_format.space_before = Pt(H1_SPACE_BEFORE_PT)

        else:
            # 正文：仿宋 小四
            text = para.text.strip()
            if _is_option_line(text):
                # 选项：左缩进 0.5cm
                para.paragraph_format.left_indent = Cm(0.5)
                para.paragraph_format.space_after = Pt(OPTION_SPACE_AFTER_PT)
            else:
                para.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
            _modify_paragraph_style(para, FONT_FANGSONG, FONT_TIMES, SIZE_SI_HAO, bold=False)

    # 调整图片宽度
    _adjust_images(doc)

    doc.save(docx_file)
    os.unlink(tmp_docx)
    for img in table_images:
        try:
            os.unlink(img)
        except OSError:
            pass
    print(f'已保存: {docx_file}')


def render_simple_table(doc, table_lines):
    if len(table_lines) < 2:
        return
    header_line = table_lines[0]
    cols_count = header_line.count('|') - 1
    if cols_count <= 0:
        return

    table = doc.add_table(rows=1, cols=cols_count)
    # 手动设置表格边框
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    header_cells = [c.strip() for c in header_line.split('|')[1:-1]]
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        # 清空默认段落
        cell._tc.clear()
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_styled_runs(p, cell_text, FONT_FANGSONG, FONT_TIMES, SIZE_XIAO_SI, italic_en=False)
        for run in p.runs:
            run.font.bold = True

    for row_line in table_lines[2:]:
        cells = [c.strip() for c in row_line.split('|')[1:-1]]
        if len(cells) == cols_count:
            row = table.add_row()
            for j, cell_text in enumerate(cells):
                cell = row.cells[j]
                cell._tc.clear()
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_styled_runs(p, cell_text, FONT_FANGSONG, FONT_TIMES, SIZE_XIAO_SI, italic_en=False)

def render_table(doc, table_content):
    lines = table_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_styled_runs(p, text, FONT_HEI, FONT_TIMES, SIZE_XIAO_ER, italic_en=False)
            for run in p.runs:
                run.font.bold = True
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_styled_runs(p, text, FONT_FZXBS, FONT_TIMES, SIZE_SAN_HAO, italic_en=False)
            for run in p.runs:
                run.font.bold = True
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            render_simple_table(doc, table_lines)
            continue
        i += 1
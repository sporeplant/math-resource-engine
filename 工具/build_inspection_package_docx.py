from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\OneDrive\math-resource-engine")
SOURCE_DIR = ROOT / "输出" / "迎检材料包_第22章"
OUTPUT = SOURCE_DIR / "人工智能赋能教育教学素养提升培训成果转化材料_第22章.docx"
APPROVED_EXAMPLE = ROOT / "输出" / "22.4_频数分布与直方图_教学设计.md"

FILES = [
    "00_封面与目录.md",
    "01_人工智能赋能教育教学素养提升培训成果转化说明.md",
    "02_第22章数字化备课成果总表.md",
    "03_第22章详案汇编.md",
    "04_第22章课件与课堂提问资源索引.md",
    "05_第22章教学日志汇编.md",
    "06_题源与教材对应说明.md",
    "07_听课笔记汇编.md",
    "08_第21章选择性佐证材料索引.md",
]

BLUE = "1F4E79"
BLUE_2 = "2F75B5"
PALE_BLUE = "D9EAF7"
PALE_GRAY = "F2F4F6"
MID_GRAY = "6B7280"
WHITE = "FFFFFF"
BLACK = "202124"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B8C4CE", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    run.font.name = "宋体"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def set_run_font(run, chinese="宋体", western="Times New Roman", size=10.5, bold=None, color=BLACK):
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), chinese)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)

    for style_name, size, color, before, after in [
        ("Heading 1", 18, BLUE, 18, 10),
        ("Heading 2", 15, BLUE, 14, 8),
        ("Heading 3", 12.5, BLUE_2, 10, 5),
        ("Heading 4", 11, BLACK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.74)
        style.paragraph_format.first_line_indent = Cm(-0.37)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.4

    if "Material Title" not in styles:
        style = styles.add_style("Material Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Material Title"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(20)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(18)
    style.paragraph_format.keep_with_next = True

    if "Metadata" not in styles:
        style = styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Metadata"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(9.5)
    style.font.color.rgb = RGBColor.from_string(MID_GRAY)
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.2


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.15)
    section.different_first_page_header_footer = True


def set_header_footer(section, label="人工智能赋能教育教学素养提升培训成果转化材料"):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label)
    set_run_font(r, chinese="微软雅黑", western="Arial", size=8.5, color=MID_GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("人工智能赋能教育教学素养提升")
    set_run_font(r, chinese="微软雅黑", western="Arial", size=16, bold=True, color=BLUE_2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("培训成果转化材料")
    set_run_font(r, chinese="微软雅黑", western="Arial", size=28, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("第22章  数据的收集、整理与描述")
    set_run_font(r, chinese="微软雅黑", western="Arial", size=17, bold=False, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(28)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:color"), BLUE_2)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    meta = [("教师", "________________"), ("学科", "初中数学"), ("年级", "八年级下册"), ("编制日期", "2026年6月10日")]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [3.0, 8.0])
    for row, (label, value) in zip(table.rows, meta):
        row.cells[0].text = label
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=120, bottom=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Cm(0)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, chinese="宋体", size=12, bold=(idx == 0))
    # Remove cover metadata borders.
    set_table_borders(table, color=WHITE, size="0")

    doc.add_page_break()


def clean_inline(text):
    text = text.replace("`", "").strip()
    replacements = {
        "content_type:": "材料类型：",
        "lesson_id:": "课时编号：",
        "lesson_name:": "课题名称：",
        "textbook_location:": "教材对应位置：",
        "scaffold_questions:": "台阶提问：",
        "student_output:": "学生预期产出：",
        "student_task:": "学生任务：",
        "teacher_action:": "教师活动：",
        "activity_type:": "活动类型：",
        "module:": "所属章节：",
        "topic:": "课题：",
        "module:": "所属章节：",
        "grade:": "适用年级：",
        "lesson_type:": "课型：",
        "duration:": "课时时长：",
        "feedback:": "反馈方式：",
        "time:": "时间：",
        "task:": "任务：",
        "aligned_objectives:": "对应学习目标：",
        "aligned_objective:": "对应学习目标：",
        "aligned_assessment:": "对应评价任务：",
        "success_criteria:": "达成标准：",
        "source_id:": "题源编号：",
        "source_type:": "题源类型：",
        "question_id:": "题目编号：",
        "source_type: textbook": "题源类型：教材",
        "source_type: exercise_bank": "题源类型：练习册题库",
    }
    for old, new in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        if text.startswith(old):
            text = new + text[len(old):]
            break
    text = text.replace("inspection_lesson", "迎检教学设计")
    heading_names = {
        "meta": "教学基本信息",
        "curriculum_alignment": "课程标准对齐",
        "knowledge_analysis": "知识分析",
        "objectives": "学习目标",
        "assessment": "评价任务",
        "lesson_flow": "教学流程",
        "practice": "课堂练习与检测",
        "homework": "分层作业",
        "answer_key": "参考答案",
        "boardwork": "板书设计",
        "consistency_matrix": "教学评一致性对照",
        "quality_check": "质量检查",
    }
    text = heading_names.get(text, text)
    for old, new in (
        ("aligned_objectives", "对应学习目标"),
        ("aligned_objective", "对应学习目标"),
        ("aligned_assessment", "对应评价任务"),
        ("success_criteria", "达成标准"),
        ("source_id", "题源编号"),
        ("source_type", "题源类型"),
        ("question_id", "题目编号"),
    ):
        text = text.replace(old, new)
    return text


def add_rich_text(paragraph, text, size=10.5, bold=False):
    text = clean_inline(text)
    parts = re.split(r"(《[^》]+》|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        emph = part.startswith("《") or (part.startswith("**") and part.endswith("**"))
        value = part[2:-2] if part.startswith("**") else part
        run = paragraph.add_run(value)
        set_run_font(run, size=size, bold=(bold or emph))


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [clean_inline(x.strip()) for x in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def widths_for_table(rows, usable_cm=16.5):
    cols = max(len(r) for r in rows)
    if cols == 2:
        return [usable_cm * 0.30, usable_cm * 0.70]
    if cols == 3:
        return [usable_cm * 0.14, usable_cm * 0.40, usable_cm * 0.46]
    if cols == 4:
        return [usable_cm * 0.09, usable_cm * 0.29, usable_cm * 0.33, usable_cm * 0.29]
    if cols == 5:
        return [usable_cm * 0.07, usable_cm * 0.19, usable_cm * 0.25, usable_cm * 0.29, usable_cm * 0.20]
    if cols == 6:
        return [usable_cm * 0.05, usable_cm * 0.15, usable_cm * 0.22, usable_cm * 0.20, usable_cm * 0.27, usable_cm * 0.11]
    return [usable_cm / cols] * cols


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, widths_for_table(rows))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    font_size = 8.2 if cols >= 5 else (8.8 if cols == 4 else 9.2)
    for ridx, values in enumerate(rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            cell.text = values[cidx] if cidx < len(values) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=95, start=105, bottom=95, end=105)
            if ridx == 0:
                set_cell_shading(cell, BLUE)
            elif ridx % 2 == 0:
                set_cell_shading(cell, PALE_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ridx == 0 or cidx == 0 or len(cell.text) <= 8) else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, chinese="宋体", size=font_size, bold=(ridx == 0), color=(WHITE if ridx == 0 else BLACK))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.first_line_indent = Cm(0)


def add_markdown_content(doc, path, skip_first_heading=False):
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    first_heading_seen = False
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = clean_inline(heading.group(2))
            if not first_heading_seen:
                first_heading_seen = True
                if skip_first_heading:
                    i += 1
                    continue
            p = doc.add_paragraph(style=f"Heading {min(level, 4)}")
            p.paragraph_format.first_line_indent = Cm(0)
            add_rich_text(p, title, bold=True)
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Cm(-0.37)
            add_rich_text(p, numbered.group(2))
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            text = clean_inline(bullet.group(1))
            # Metadata-like fields are visually quieter than teaching content.
            if re.match(r"[A-Za-z_]+:\s*", text) or re.match(r"(材料类型|课时编号|课题名称|所属章节|适用年级|课型|课时时长|对应学习目标|对应评价任务|达成标准|题源编号|题源类型|题目编号|时间|教材对应位置|活动类型|学生任务|教师活动|台阶提问|预期产出|授课日期|授课班级|对应教学设计|实施记录|课堂证据|改进动作|活动名称|主讲教师|活动日期|活动地点|记录主题|授课教师|听课日期|年级班级|课题|整理依据)[：:]", text):
                p = doc.add_paragraph(style="Metadata")
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.first_line_indent = Cm(0)
                normalized = text.replace("：", ":", 1)
                label, sep, value = normalized.partition(":")
                r = p.add_run(label + ("：" if sep else ""))
                set_run_font(r, chinese="微软雅黑", western="Arial", size=9.5, bold=True, color=BLUE)
                if value:
                    r = p.add_run(value.strip())
                    set_run_font(r, size=9.5, color=MID_GRAY)
            else:
                p = doc.add_paragraph(style="List Bullet")
                add_rich_text(p, text)
            i += 1
            continue
        # Merge adjacent prose lines into one paragraph.
        block = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|") or re.match(r"^(\d+\.\s+|[-*]\s+)", nxt):
                break
            block.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_rich_text(p, " ".join(block))


def add_contents_page(doc):
    p = doc.add_paragraph(style="Material Title")
    p.add_run("目  录")
    entries = [
        ("材料一", "人工智能赋能教育教学素养提升培训成果转化说明"),
        ("材料二", "第22章数字化备课成果总表"),
        ("材料三", "第22章已审核示范课例"),
        ("材料四", "第22章课件与课堂提问资源索引"),
        ("材料五", "第22章教学日志汇编"),
        ("材料六", "题源与教材对应说明"),
        ("材料七", "听课笔记汇编"),
        ("材料八", "第21章选择性佐证材料索引"),
    ]
    for label, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.right_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(label + "  ")
        set_run_font(r, chinese="微软雅黑", western="Arial", size=11, bold=True, color=BLUE)
        r = p.add_run(title)
        set_run_font(r, size=11)
    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(0)
    r = note.add_run("说明：各项材料均按装订顺序另页起排，便于查阅、抽检和归档。")
    set_run_font(r, size=9.5, color=MID_GRAY)
    doc.add_page_break()


def add_approved_example(doc):
    lines = APPROVED_EXAMPLE.read_text(encoding="utf-8").splitlines()
    # Strip machine front matter and replace it with a concise print-facing audit record.
    second_rule = None
    for idx in range(1, len(lines)):
        if lines[idx].strip() == "---":
            second_rule = idx
            break
    body = lines[second_rule + 1:] if second_rule is not None else lines
    temp = SOURCE_DIR / "_approved_example_tmp.md"
    audit = [
        "# 22.4 频数分布与直方图教学设计",
        "",
        "## 审核信息",
        "",
        "- 课时编号：22.4",
        "- 课型：新授课",
        "- 课时时长：40分钟",
        "- 审核状态：教师审核通过",
        "- 审核日期：2026年6月9日",
        "- 审核意见：表图转化环节采用阶梯电价同一组数据进行小台阶作图，教材A组第2题后移为机动迁移任务。",
        "",
    ]
    temp.write_text("\n".join(audit + body), encoding="utf-8")
    add_markdown_content(doc, temp, skip_first_heading=False)
    temp.unlink(missing_ok=True)


def main():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    set_header_footer(doc.sections[0])
    add_cover(doc)

    add_contents_page(doc)

    # Use 00 as front matter, but omit its duplicate cover heading and directory table.
    front = SOURCE_DIR / FILES[0]
    text = front.read_text(encoding="utf-8")
    start = text.find("## 一、材料说明")
    end = text.find("## 三、装订目录")
    temp = ROOT / "输出" / "迎检材料包_第22章" / "_frontmatter_tmp.md"
    front_text = "# 编印说明\n\n" + text[start:end].replace("## 一、材料说明", "## 材料说明").replace("## 二、阅读导航", "## 查阅建议")
    tail = text[text.find("## 四、主要教学思路"):]
    front_text += tail
    temp.write_text(front_text, encoding="utf-8")
    p = doc.add_paragraph(style="Material Title")
    p.add_run("编 印 说 明")
    add_markdown_content(doc, temp, skip_first_heading=True)
    temp.unlink(missing_ok=True)

    material_labels = ["材料一", "材料二", "材料三", "材料四", "材料五", "材料六", "材料七", "材料八"]
    for label, filename in zip(material_labels, FILES[1:]):
        doc.add_page_break()
        path = SOURCE_DIR / filename
        title = path.read_text(encoding="utf-8").splitlines()[0].lstrip("# ").strip()
        kicker = doc.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kicker.paragraph_format.first_line_indent = Cm(0)
        kicker.paragraph_format.space_after = Pt(5)
        r = kicker.add_run(label)
        set_run_font(r, chinese="微软雅黑", western="Arial", size=10, bold=True, color=BLUE_2)
        p = doc.add_paragraph(style="Material Title")
        p.add_run(title)
        add_markdown_content(doc, path, skip_first_heading=True)
        if filename == "03_第22章详案汇编.md":
            doc.add_page_break()
            add_approved_example(doc)

    # Core properties and print settings.
    doc.core_properties.title = "人工智能赋能教育教学素养提升培训成果转化材料（第22章）"
    doc.core_properties.subject = "初中数学数字化备课成果汇编"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "人工智能赋能教育；初中数学；教学设计；成果转化"
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
